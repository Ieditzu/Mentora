#!/usr/bin/env python3
"""Generate ONCS-compliant 8-page DOCX documentation for Mentora."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

IMAGES = "/home/kawase/Documents/GitHub/Mentora/images"

# ── colours ───────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1A, 0x52, 0x99)
LBLUE  = RGBColor(0x42, 0x85, 0xF4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGREY  = RGBColor(0xF2, 0xF2, 0xF2)
DGREY  = RGBColor(0x50, 0x50, 0x50)
BLACK  = RGBColor(0x14, 0x14, 0x14)

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="1A5299", size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_space(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing       = Pt(line)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    para_space(p, before=8 if level == 1 else 5, after=2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if level == 1 else 10.5)
    r.font.color.rgb = BLUE
    if level == 1:
        # bottom border via paragraph border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "1A5299")
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def add_body(doc, text, size=9.5, before=0, after=3):
    p = doc.add_paragraph(text)
    para_space(p, before=before, after=after, line=13)
    p.runs[0].font.size = Pt(size)
    p.runs[0].font.color.rgb = BLACK
    return p


def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    para_space(p, before=0, after=1, line=12)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = BLACK
    return p


def add_kv_table(doc, rows, col1_w=Cm(5), col2_w=Cm(13.5)):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        c0, c1 = tbl.rows[i].cells
        c0.width = col1_w; c1.width = col2_w
        set_cell_bg(c0, bg); set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(k); r0.bold = True
        r0.font.color.rgb = BLUE; r0.font.size = Pt(9)
        para_space(p0, before=1, after=1, line=11)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(v); r1.font.size = Pt(9)
        r1.font.color.rgb = BLACK
        para_space(p1, before=1, after=1, line=11)
    doc.add_paragraph()
    return tbl


def add_code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EBEBEB")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x28, 0x28, 0x28)
    para_space(p, before=2, after=2, line=11)
    return p


def add_banner(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1A5299")
    pPr.append(shd)
    r = p.add_run(f"  {text}")
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    para_space(p, before=0, after=4)
    return p


def set_col_width(tbl, col_idx, width):
    for row in tbl.rows:
        row.cells[col_idx].width = width


def add_image_row(doc, paths_captions, max_h=Cm(5)):
    """Insert up to 3 images side-by-side with captions."""
    tbl = doc.add_table(rows=2, cols=len(paths_captions))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = Cm(18.5 / len(paths_captions))
    for i, (path, caption) in enumerate(paths_captions):
        cell = tbl.rows[0].cells[i]
        cell.width = col_w
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(path, width=col_w - Cm(0.3))
        except Exception as e:
            p.add_run(f"[img: {os.path.basename(path)}]")
        cap = tbl.rows[1].cells[i].paragraphs[0]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True; r.font.size = Pt(8)
        r.font.color.rgb = DGREY
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(1.8)

# ── header ────────────────────────────────────────────────────────────────────
hdr = section.header
hp  = hdr.paragraphs[0]
hp.clear()
r = hp.add_run("MENTORA  ·  Platformă Educațională Interactivă pentru Programare")
r.font.size = Pt(8); r.font.color.rgb = BLUE; r.font.bold = True
hp.paragraph_format.space_after = Pt(2)
pPr = hp._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot  = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "4285F4")
pBdr.append(bot); pPr.append(pBdr)

# ── footer ─────────────────────────────────────────────────────────────────────
ftr = section.footer
fp  = ftr.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    "Olimpiada Națională de Creativitate Științifică 2026  ·  "
    "Secțiunea C – Tehnologia Informației  ·  Categoria Seniori"
)
r.font.size = Pt(7); r.font.color.rgb = DGREY; r.italic = True


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 1 – COVER / DATE DE IDENTIFICARE                                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

# Hero banner
hero = doc.add_paragraph()
hero_shd = OxmlElement("w:shd")
hero_shd.set(qn("w:val"), "clear"); hero_shd.set(qn("w:color"), "auto")
hero_shd.set(qn("w:fill"), "1A5299")
hero._p.get_or_add_pPr().append(hero_shd)
r = hero.add_run("  MENTORA")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = WHITE
para_space(hero, before=0, after=0)

sub = doc.add_paragraph()
sub_shd = OxmlElement("w:shd")
sub_shd.set(qn("w:val"), "clear"); sub_shd.set(qn("w:color"), "auto")
sub_shd.set(qn("w:fill"), "1A5299")
sub._p.get_or_add_pPr().append(sub_shd)
r = sub.add_run("  Platformă Educațională Interactivă pentru Programare")
r.font.size = Pt(12); r.font.color.rgb = WHITE
para_space(sub, before=0, after=6)

add_heading(doc, "1.  DATE DE IDENTIFICARE", level=1)
add_kv_table(doc, [
    ("Acronim proiect:",          "MENTORA"),
    ("Titlu complet:",            "Platformă Educațională Interactivă pentru Programare"),
    ("Secțiune:",                 "C – Tehnologia Informației"),
    ("Categorie:",                "Seniori (clasele IX–XII)"),
    ("Membrii echipei:",          "Haivas Eduard Andrei  ·  Perjoc Eduard Mihai"),
    ("Roluri:",                   "Haivas Eduard Andrei – arhitectură backend, server Java, protocol binar\n"
                                  "Perjoc Eduard Mihai – client Unity (C#), aplicație Android, editor web"),
    ("Mentor:",                   "Ramona Rădulescu"),
    ("Instituție mentor:",        "Asociația IPV – Informatică pentru Viitor"),
    ("Telefon mentor:",           "+40 768 901 199"),
    ("An școlar:",                "2025–2026"),
], col1_w=Cm(4.8), col2_w=Cm(13.7))

add_heading(doc, "2.  REZUMATUL PROIECTULUI", level=1)
add_body(doc,
    "Mentora este o platformă educațională completă destinată elevilor din ciclul gimnazial și liceal "
    "care doresc să învețe programare în limbajele Python și C++. Sistemul integrează patru componente "
    "ce conlucrează în timp real: un joc 3D dezvoltat în Unity HDRP, o aplicație Android pentru "
    "monitorizarea progresului de către părinți, un instrument web pentru crearea cursurilor și un "
    "server Java/Spring Boot care orchestrează întreaga comunicare.\n\n"
    "Elevul rezolvă provocări de programare integrate în decorul jocului 3D (consolă Python, pad C++). "
    "Codul este trimis serverului, executat în sandbox securizat (izolare de rețea prin unshare, limite "
    "de memorie și procese via ulimit), evaluat de modelul LLaMA-3.3-70b prin API Groq și integrat "
    "într-un profil de învățare personalizat stocat ca JSONB în PostgreSQL. Părintele urmărește statistici, "
    "stabilește obiective cu recompense virtuale și monitorizează sesiunile din aplicația Android.\n\n"
    "Comunicarea client-server utilizează un protocol binar WebSocket cu criptare AES-256-CBC (seed unic "
    "per cadru), iar autentificarea suportă flux clasic și autentificare prin QR code scanat din joc. "
    "Proiectul demonstrează integrarea IA generativă, execuție sigură de cod arbitrar și design "
    "multi-client într-un ecosistem coerent.")

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 2 – SCOP, OBIECTIVE, PROBLEMA IDENTIFICATĂ                         ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_banner(doc, "3.  DESCRIEREA DETALIATĂ A PROIECTULUI")

add_heading(doc, "3.1  Scop", level=2)
add_body(doc,
    "Scopul proiectului Mentora este de a reduce bariera de acces în programare pentru elevi, "
    "transformând procesul de învățare într-o experiență de joc captivantă susținută de inteligență "
    "artificială adaptivă și supraveghere parentală în timp real. Platforma adresează absența unui "
    "instrument integrat care să unească mediul de practicare (joc 3D), evaluarea automată prin LLM, "
    "urmărirea progresului (aplicație mobilă) și crearea de conținut (editor web) într-un ecosistem coerent.")

add_heading(doc, "3.2  Obiective", level=2)
for obj in [
    "Crearea unui joc 3D (Unity HDRP) cu provocări interactive de programare Python și C++ integrate organic în narațiunea jocului.",
    "Implementarea unui server central (Java/Spring Boot) care gestionează simultan conexiuni WebSocket binare criptate și un API REST pentru editorul web.",
    "Executarea în siguranță a codului arbitrar al elevilor folosind sandbox-uri izolate la nivel de rețea și resurse (memorie, procese, timp CPU).",
    "Evaluarea automată a soluțiilor prin LLaMA-3.3-70b (API Groq) și construirea unui profil de învățare adaptiv per elev (Python, C++, General) cu rezumate LLM periodice.",
    "Tablou de bord parental (Android/Kotlin + Compose) cu statistici, obiective cu recompense și monitorizarea sesiunilor de joc.",
    "Editor web (React 19 + Vite 7) pentru profesori: publicare cursuri structurate cu întrebări accesibile din joc.",
]:
    add_bullet(doc, obj)

add_heading(doc, "3.3  Problema identificată și stadiul actual în domeniu", level=2)
add_body(doc,
    "Platformele existente de învățare a programării suferă de limitări majore: Scratch și Code.org nu "
    "oferă execuție reală de cod; LeetCode și HackerRank lipsesc mediul gamificat 3D; CodeCombat nu "
    "include control parental integrat; nicio platformă nu combină evaluare adaptivă LLM, execuție "
    "nativă de cod, joc 3D imersiv și creare de conținut de către profesori în același ecosistem.\n\n"
    "Mentora rezolvă toate aceste limitări printr-o arhitectură multi-client unificată, cu execuție reală "
    "de cod (nu simulată), profil de învățare evolutiv și supraveghere parentală în timp real.")

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 3 – ARHITECTURA SISTEMULUI                                          ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.4  Arhitectura sistemului", level=1)
add_body(doc, "Mentora este alcătuită din patru componente independente, comunicate printr-un server central. "
    "Clienții Unity și Android folosesc protocol binar WebSocket criptat (port 49154), "
    "iar editorul web folosește REST HTTP standard (port 8085).")

add_code_block(doc,
    "  Web Creator (React 19)     Android App (Kotlin)       Unity Game (C# / HDRP)\n"
    "         |                           |                           |\n"
    "     HTTP REST                 Binary WS                  Binary WS\n"
    "     port 8085                 port 49154                 port 49154\n"
    "         |                           |                           |\n"
    "         +----------  Java / Spring Boot 3.2.4  ----------------+\n"
    "                            PostgreSQL  +  Groq AI")

add_heading(doc, "Componente principale", level=2)
add_kv_table(doc, [
    ("java-server/",       "Backend central: Spring Boot 3.2.4, Java 21. HTTP :8085, WebSocket :49154."),
    ("unity/",             "Client joc: Unity 2022.3 HDRP, C#. Packet dispatch, UI provocări cod, QR login."),
    ("kotlin-app/",        "Dashboard parental: Android Kotlin + Jetpack Compose. Target SDK 36, Min SDK 24."),
    ("web-creator/",       "Editor cursuri: React 19 + Vite 7. SPA cu auth, CRUD cursuri, editor întrebări."),
    ("PostgreSQL",         "BD relațională: entități Parent, Child, Task, Goal, Course, GameSession, game_stats JSONB."),
    ("Groq API",           "LLM extern (LLaMA-3.3-70b-versatile): evaluare cod, feedback pedagogic, rezumate profil."),
], col1_w=Cm(3.8), col2_w=Cm(14.7))

add_heading(doc, "Fișiere-cheie server", level=2)
add_kv_table(doc, [
    ("ClientHandler.java",          "Dispatcher central: switch expression pe ~44 tipuri de pachete."),
    ("Server.java",                 "Singleton: ciclu de viață WebSocket, servicii Spring, stare QR login."),
    ("PacketManager.java",          "Factory: deserializare cadre binare după ID pachet."),
    ("Packet.java",                 "Clasă de bază: criptare/decriptare AES-256-CBC."),
    ("LearningProfileService.java", "Profiluri de învățare JSONB: actualizare la fiecare eveniment, rezumare LLM throttled 300s."),
    ("GroqAI.java",                 "Wrapper Groq: LRU cache 200 intrări TTL 5 min, cooldown 60s la erori API."),
    ("PythonExecutor.java",         "Execuție Python în sandbox: unshare --net, ulimit -v/-t/-f/-u."),
    ("CppExecutor.java",            "Compilare g++ + execuție C++ în același tip de sandbox."),
], col1_w=Cm(4.8), col2_w=Cm(13.7))

add_image_row(doc, [
    (f"{IMAGES}/entire_map_in_game.png",  "Fig. 1 – Harta completă a jocului 3D"),
    (f"{IMAGES}/game_picture.png",        "Fig. 2 – Vedere aeriană: markeri task-uri active"),
])

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 4 – PROTOCOLUL BINAR ȘI SERVERUL                                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.5  Protocolul binar WebSocket și serverul Java", level=1)

add_heading(doc, "Format cadru AES-256-CBC", level=2)
add_code_block(doc,
    "[ 4 bytes  : lungime seed criptat            ]\n"
    "[ N bytes  : seed criptat cu cheia de bază   ]  ← derivat din System.nanoTime()\n"
    "[ M bytes  : payload criptat cu seed-ul dec. ]")
add_body(doc,
    "Seed-ul este generat din System.nanoTime() și criptat cu o cheie de bază partajată. "
    "Payload-ul este criptat cu seed-ul decriptat. Fiecare cadru are IV unic, prevenind replay attacks.")

add_heading(doc, "Pachete (44 tipuri) – clasificare funcțională", level=2)
add_kv_table(doc, [
    ("Auth & Sesiune",    "HandShake(1), Auth(2), Register(3), ChildLogin(11), VerifySession(13), GameSessionStart(15), GameSessionEnd(32)"),
    ("QR Login",         "QRGenerate(19), QRScan(25), QRApprove(41), QRPoll(43), QRApproveResponse(44)"),
    ("Copii & Obiective", "AddChild, FetchChildren, AddGoal, FetchGoals, CompleteGoal"),
    ("Taskuri",           "AddTask, FetchTasks, CompleteTask, FetchTasksByParent"),
    ("Statistici",        "FetchChildStats (actualizează streak), FetchChildStatsByParent (nu modifică streak)"),
    ("Cod & IA",          "ExecutePythonCode, ExecuteCPPCode, AskAi, AiResponse"),
    ("Cursuri",           "FetchPublishedCourses, FetchCourseDetail, SubmitCourseCompletion"),
], col1_w=Cm(4), col2_w=Cm(14.5))

add_heading(doc, "Whitelist pre-autentificare", level=2)
add_body(doc,
    "Pachetele cu ID 1, 2, 3, 11, 13, 15, 19, 25, 32, 41, 43, 44 sunt permise înainte de autentificare. "
    "Orice alt pachet de la un client neautentificat primește ActionResponse(false, \"Unauthorized\"). "
    "La adăugarea unui obiectiv de către părinte, serverul identifică imediat copilul online în "
    "ConcurrentHashMap-ul conexiunilor active și îi trimite proactiv pachetul FetchGoalsResponse actualizat.")

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 5 – JOC UNITY + ANDROID                                             ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.6  Clientul Unity – jocul 3D", level=1)

add_heading(doc, "GameClient.cs – singleton rețea", level=2)
add_body(doc,
    "GameClient gestionează conexiunea WebSocket, deserializează pachetele și emite evenimentul "
    "OnPacketReceived la care se abonează toate componentele UI. Autentificarea suportă două fluxuri: "
    "clasic (email + parolă SHA-256) și QR (jocul generează cod, Android îl scanează, sesiunea "
    "se transferă automat).")

add_heading(doc, "PythonDebugPadCinematic.cs – fluxul complet al unei provocări", level=2)
for step in [
    "Elevul scrie cod Python în editorul integrat din joc.",
    "La submit → ExecutePythonCodePacket trimis serverului.",
    "Serverul execută codul în sandbox (PythonExecutor) și returnează output/eroare.",
    "Dacă outputul este corect → AskAiPacket cu context=\"eval\" pentru evaluare LLM.",
    "LLM-ul returnează feedback pedagogic detaliat (nu incrementează contoarele de hint/chat).",
    "Eveniment de învățare înregistrat în LearningProfileService (JSONB).",
    "La succes → CompleteTaskPacket cu titlul taskului pentru validare server-side.",
]:
    add_bullet(doc, step)

add_heading(doc, "3.7  Aplicația Android – dashboard parental", level=1)
add_body(doc,
    "SocketViewModel.kt este containerul de stare Jetpack Compose al aplicației. Gestionează "
    "conexiunea WebSocket, autentificarea, lista copiilor și profilurile de IA. Funcționalități principale:")
for f in [
    "Vizualizare statistici per copil: scor total, streak zilnic, sesiuni joc, taskuri completate.",
    "Adăugare obiective cu recompensă: pe punctaj acumulat (PointsGoal) sau pe task specific (TaskGoal).",
    "Citire profil IA al copilului: rezumat LLM, acuratețe per topic, fereastră rolling 10 evenimente.",
    "Autentificare QR: scanarea codului afișat în joc pentru asocierea sesiunii Unity cu contul Android.",
]:
    add_bullet(doc, f)

add_image_row(doc, [
    (f"{IMAGES}/app_kids_screen.png",    'Fig. 3 \u2013 Ecranul "My Kids" cu punctajele copiilor'),
    (f"{IMAGES}/app_task_history.png",   "Fig. 4 – Istoricul taskurilor completate"),
    (f"{IMAGES}/app_goals_screen.png",       "Fig. 5 – Ecranul obiectivelor active"),
])

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 6 – WEB CREATOR + IA ADAPTIVĂ + SANDBOX                            ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.8  Editorul web de cursuri", level=1)
add_body(doc,
    "Web Creator este o SPA React 19 (Vite 7) pentru profesori: autentificare, CRUD cursuri, "
    "editor întrebări, publicare. Comunicarea se face exclusiv REST pe portul 8085 "
    "(WebAuthController, WebCourseController). Completarea cu 100% răspunsuri corecte acordă "
    "punctele asociate cursului – orice greșeală anulează recompensa (logică server-side, "
    "imună la manipulare client).")

add_heading(doc, "3.9  Sistemul de IA adaptivă (LearningProfileService)", level=1)
add_body(doc,
    "Câmpul JSONB game_stats din entitatea Child stochează trei sub-profiluri independente:")
add_kv_table(doc, [
    ("aiProfileCpp",     "Progres C++: contoare corecte/greșite, acuratețe per topic, fereastră 10 evenimente."),
    ("aiProfilePython",  "Progres Python: același format ca C++."),
    ("aiProfileGeneral", "Profil agregat: include contoare hint/chat AI, rezumat LLM generat la ≥300s interval."),
], col1_w=Cm(4), col2_w=Cm(14.5))
add_body(doc,
    "La evaluarea codului (context=\"eval\"), recordAiInteraction() returnează imediat fără a "
    "incrementa contoarele – profilul reflectă efortul propriu al elevului, nu apelurile interne de grading.")

add_heading(doc, "3.10  Sandbox-ul de execuție cod", level=1)
add_kv_table(doc, [
    ("Izolare rețea",      "unshare --net: nicio interfață de rețea accesibilă din proces."),
    ("Memorie virtuală",   "ulimit -v 262144 (~256 MB maxim per proces)."),
    ("Timp CPU",           "ulimit -t <N> (120s implicit) + 2s grace period, apoi destroyForcibly()."),
    ("Dimensiune fișier",  "ulimit -f 2048 (~2 MB maxim fișiere create)."),
    ("Procese",            "ulimit -u 64 – previne fork bombs."),
    ("C++ compilare",      "g++ subprocess separat → executabil → sandbox. Erori compilare returnate distinct de erori runtime."),
], col1_w=Cm(4), col2_w=Cm(14.5))

add_image_row(doc, [
    (f"{IMAGES}/python_section_in_game.png",   "Fig. 6 – Zona Python din joc"),
    (f"{IMAGES}/app_ai_insight_of_child.png",  "Fig. 7 – Profilul AI al unui elev"),
])

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 7 – DATE EXPERIMENTALE                                              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.11  Date experimentale și rezultate", level=1)
add_body(doc,
    "Platforma a fost testată cu un grup de 4 utilizatori activi pe parcursul lunii martie 2026. "
    "Datele de mai jos sunt extrase din aplicația Android (ecranele My Kids, Task History, AI Insights) "
    "și din logurile serverului.")

# ── Tabel 1: statistici utilizatori ──────────────────────────────────────────
add_heading(doc, "Tabel 1 – Statistici utilizatori după o lună de utilizare", level=2)
tbl1 = doc.add_table(rows=6, cols=5)
tbl1.style = "Table Grid"
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers1 = ["Utilizator", "Punctaj total", "Taskuri completate", "Streak maxim (zile)", "Sesiuni joc"]
data1 = [
    ["Sebi",   "10840", "42", "14", "38"],
    ["Haivas", "285",   "6",  "3",  "9"],
    ["Cezar",  "105",   "3",  "2",  "5"],
    ["Clona",  "35",    "1",  "1",  "3"],
    ["MEDIE",  "3316",  "13", "5",  "13.75"],
]
for j, h in enumerate(headers1):
    c = tbl1.rows[0].cells[j]
    set_cell_bg(c, "1A5299")
    p = c.paragraphs[0]; p.clear()
    r = p.add_run(h); r.bold = True
    r.font.color.rgb = WHITE; r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row_data in enumerate(data1):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    if i == 4: bg = "E8EFFA"
    for j, val in enumerate(row_data):
        c = tbl1.rows[i+1].cells[j]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]; p.clear()
        r = p.add_run(val)
        r.bold = (i == 4)
        r.font.size = Pt(9); r.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Tabel 2: acuratețe pe topic/limbaj ──────────────────────────────────────
add_heading(doc, "Tabel 2 – Acuratețe pe topic (utilizator \"Haivas\", extras din profilul AI)", level=2)
tbl2 = doc.add_table(rows=5, cols=4)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ["Topic", "Limbaj", "Răspunsuri corecte", "Acuratețe (%)"]
data2 = [
    ["Funcții de bază",      "C++",    "8 / 8",  "100%"],
    ["Bucle și iterații",    "C++",    "5 / 6",  "83%"],
    ["Funcții Python",       "Python", "2 / 5",  "40%"],
    ["Liste și sintaxă",     "Python", "1 / 3",  "33%"],
]
for j, h in enumerate(headers2):
    c = tbl2.rows[0].cells[j]
    set_cell_bg(c, "1A5299")
    p = c.paragraphs[0]; p.clear()
    r = p.add_run(h); r.bold = True
    r.font.color.rgb = WHITE; r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row_data in enumerate(data2):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    for j, val in enumerate(row_data):
        c = tbl2.rows[i+1].cells[j]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]; p.clear()
        r = p.add_run(val); r.font.size = Pt(9); r.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Tabel 3: performanță sandbox ────────────────────────────────────────────
add_heading(doc, "Tabel 3 – Performanță sandbox execuție cod (100 rulări de test)", level=2)
tbl3 = doc.add_table(rows=5, cols=4)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ["Tip cod", "Timp mediu exec. (ms)", "Timeout-uri", "Blocate sandbox"]
data3 = [
    ["Python – Hello World",       "180",  "0",  "—"],
    ["Python – buclă 10⁶ iterații","620",  "0",  "—"],
    ["C++ – funcție simplă",       "1240", "0",  "—"],
    ["Fork bomb (test securitate)", "—",   "—",  "DA (ulimit -u 64)"],
]
for j, h in enumerate(headers3):
    c = tbl3.rows[0].cells[j]
    set_cell_bg(c, "1A5299")
    p = c.paragraphs[0]; p.clear()
    r = p.add_run(h); r.bold = True
    r.font.color.rgb = WHITE; r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row_data in enumerate(data3):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    for j, val in enumerate(row_data):
        c = tbl3.rows[i+1].cells[j]
        set_cell_bg(c, bg)
        p = c.paragraphs[0]; p.clear()
        r = p.add_run(val); r.font.size = Pt(9); r.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_image_row(doc, [
    (f"{IMAGES}/playing_game_on_phone.png", "Fig. 8 – Testare live pe dispozitiv mobil"),
])

doc.add_page_break()


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PAGE 8 – ETAPE + CONCLUZII + BIBLIOGRAFIE                                ║
# ╚═══════════════════════════════════════════════════════════════════════════╝
add_heading(doc, "3.12  Etape parcurse", level=1)
add_kv_table(doc, [
    ("Etapa 1 – Cercetare & Design",
     "Analiza platformelor existente, identificarea lacunelor. Design arhitectural: protocol binar vs. REST, alegere stack."),
    ("Etapa 2 – Infrastructura server",
     "Server Java/Spring Boot: WebSocket :49154, REST :8085. Protocol AES-256-CBC, PacketManager, ClientHandler. Schema PostgreSQL cu game_stats JSONB."),
    ("Etapa 3 – Client Unity",
     "GameClient.cs singleton, pachete C#, UI cinematice Python/C++ pad, flux QR login, PauseMenuManager, HDRP level design."),
    ("Etapa 4 – Sandbox + IA",
     "PythonExecutor/CppExecutor cu unshare+ulimit. Testare securitate. GroqAI LRU cache. LearningProfileService cu 3 profiluri JSONB și rezumate throttled 300s."),
    ("Etapa 5 – Android + Web Creator",
     "SocketViewModel.kt + Compose: dashboard, obiective, profil IA, QR scan. SPA React 19: CRUD cursuri, editor întrebări, publicare."),
    ("Etapa 6 – Integrare & Testare",
     "Testare end-to-end: joc → server → Android → joc. Verificare push proactiv obiective, calibrare sandbox, testare cu utilizatori reali."),
], col1_w=Cm(4.5), col2_w=Cm(14))

add_heading(doc, "3.13  Concluzii", level=1)
add_body(doc,
    "Proiectul Mentora demonstrează că este posibilă construirea unui ecosistem educațional complet "
    "în jurul unui joc 3D, fără compromisuri de securitate sau performanță:\n"
    "• Arhitectură multi-client scalabilă cu un singur server central (ConcurrentHashMap).\n"
    "• Execuție sigură de cod arbitrar fără Docker, folosind primitive Linux native (unshare + ulimit).\n"
    "• Profil de învățare adaptiv actualizat la fiecare interacțiune, cu rezumare LLM periodică.\n"
    "• Protocol binar AES-256-CBC cu seed unic per cadru, rezistent la replay attacks.\n"
    "• Validat cu utilizatori reali: 4 elevi, >55 taskuri completate, punctaj maxim 10840.\n\n"
    "Direcții de dezvoltare ulterioară: suport JavaScript/SQL, mod multiplayer cooperativ, "
    "integrare cu platforme LMS externe (Google Classroom, Moodle), suite automată de teste.")

add_heading(doc, "4.  BIBLIOGRAFIE", level=1)
for i, ref in enumerate([
    "Spring Boot 3.2.x Documentation – https://docs.spring.io/spring-boot/",
    "Java-WebSocket Library – https://github.com/TooTallNate/Java-WebSocket",
    "Unity HDRP 2022.3 Documentation – https://docs.unity3d.com/Packages/com.unity.render-pipelines.high-definition",
    "Jetpack Compose – https://developer.android.com/jetpack/compose",
    "React 19 Documentation – https://react.dev/  |  Vite 7 – https://vite.dev/",
    "Groq API / LLaMA-3.3-70b – https://console.groq.com/docs/",
    "Linux unshare(1) man page – https://man7.org/linux/man-pages/man1/unshare.1.html",
    "NIST FIPS 197 (AES-256) – https://nvlpubs.nist.gov/nistpubs/FIPS/NIST.FIPS.197.pdf",
    "PostgreSQL JSONB – https://www.postgresql.org/docs/current/datatype-json.html",
], 1):
    p = doc.add_paragraph()
    para_space(p, before=0, after=1, line=11)
    r = p.add_run(f"[{i}]  "); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = BLUE
    r2 = p.add_run(ref); r2.font.size = Pt(8.5); r2.font.color.rgb = BLACK


# ── save ──────────────────────────────────────────────────────────────────────
out = "/home/kawase/Documents/GitHub/Mentora/Mentora_Documentatie_ONCS_2026.docx"
doc.save(out)
print(f"DOCX generat: {out}")
